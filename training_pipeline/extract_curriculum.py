#!/usr/bin/env python3
"""
extract_curriculum.py
Reads NaCCA curriculum PDFs and DOCX files from curriculum_upload/
and extracts structured data to update nacca_db.json
"""

import os
import json
import sys
from pathlib import Path

ROOT = Path(__file__).parent.parent

def extract_from_docx(filepath):
    """Extract text from DOCX file."""
    try:
        from docx import Document
        doc = Document(filepath)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return '\n'.join(text)
    except Exception as e:
        print(f"  ⚠️  DOCX read error for {filepath.name}: {e}")
        return ''

def extract_from_pdf(filepath):
    """Extract text from PDF file."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(filepath))
        text = []
        for page in doc:
            text.append(page.get_text())
        doc.close()
        return '\n'.join(text)
    except Exception as e:
        try:
            import pdf2text
            return pdf2text.extract(str(filepath))
        except:
            print(f"  ⚠️  PDF read error for {filepath.name}: {e}")
            print(f"      Install: pip install pymupdf")
            return ''

def process_class_folder(class_code, folder_path):
    """Process all files in a class folder."""
    results = []
    folder = Path(folder_path)
    if not folder.exists():
        return results

    files = list(folder.glob('*.pdf')) + list(folder.glob('*.docx')) + list(folder.glob('*.doc'))
    if not files:
        print(f"  📂 {class_code}: No files found (add curriculum PDFs/DOCX here)")
        return results

    for f in files:
        print(f"  Processing: {f.name}")
        if f.suffix.lower() == '.pdf':
            text = extract_from_pdf(f)
        else:
            text = extract_from_docx(f)

        if text:
            results.append({
                'source': f.name,
                'class': class_code,
                'text_length': len(text),
                'text_preview': text[:500],
                'full_text': text
            })
            print(f"     Extracted {len(text):,} characters")

    return results

def main():
    print("\n" + "="*60)
    print("  LessonGen Ghana — Curriculum Extractor")
    print("="*60 + "\n")

    upload_root = ROOT / 'curriculum_upload'
    output_file = ROOT / 'training_pipeline' / 'extracted_curriculum.json'
    report_file = ROOT / 'training_pipeline' / 'extraction_report.txt'

    if not upload_root.exists():
        print(f"❌ curriculum_upload/ folder not found at {upload_root}")
        sys.exit(1)

    class_map = {
        'KG/KG1': 'KG1', 'KG/KG2': 'KG2',
        'Primary/B1': 'B1', 'Primary/B2': 'B2', 'Primary/B3': 'B3',
        'Primary/B4': 'B4', 'Primary/B5': 'B5', 'Primary/B6': 'B6',
        'JHS/B7': 'B7', 'JHS/B8': 'B8', 'JHS/B9': 'B9',
    }

    all_results = {}
    total_files = 0
    total_chars = 0
    empty_folders = []

    for rel_path, class_code in class_map.items():
        folder = upload_root / rel_path
        print(f"\n[Level] {class_code} -- {folder}")
        results = process_class_folder(class_code, folder)
        if results:
            all_results[class_code] = results
            total_files += len(results)
            total_chars += sum(r['text_length'] for r in results)
        else:
            empty_folders.append(class_code)

    # Save extracted data
    output_file.parent.mkdir(parents=True, exist_ok=True)
    text_dir = output_file.parent / 'extracted_text'
    text_dir.mkdir(exist_ok=True)

    save_data = {}
    for k, v in all_results.items():
        save_data[k] = []
        for r in v:
            # Save full text to a separate file
            txt_filename = f"{r['class']}_{r['source'].replace('.pdf', '.txt').replace('.docx', '.txt')}"
            txt_path = text_dir / txt_filename
            with open(txt_path, 'w', encoding='utf-8') as tf:
                tf.write(r['full_text'])
            
            save_data[k].append({
                'source': r['source'],
                'class': r['class'],
                'text_length': r['text_length'],
                'text_preview': r['text_preview'],
                'text_file': txt_filename
            })

    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(save_data, f, indent=2, ensure_ascii=False)

    # Write report
    report = []
    report.append("CURRICULUM EXTRACTION REPORT")
    report.append("="*50)
    report.append(f"Total files processed: {total_files}")
    report.append(f"Total text extracted: {total_chars:,} characters")
    report.append(f"\nFolders with content: {list(all_results.keys())}")
    report.append(f"\nEmpty folders (add curriculum files here):")
    for ef in empty_folders:
        report.append(f"  - curriculum_upload/ ... /{ef}/")
    report.append("\nNext step: run python training_pipeline/build_training_data.py")

    with open(report_file, 'w', encoding='utf-8') as f:
        f.write('\n'.join(report))

    print(f"\n{'='*60}")
    print(f"Extraction complete!")
    print(f"   Files: {total_files} | Characters: {total_chars:,}")
    print(f"   Output: {output_file}")
    print(f"   Report: {report_file}")
    if empty_folders:
        print(f"\nWarning: Empty folders (no curriculum files yet):")
        for ef in empty_folders:
            print(f"   curriculum_upload/.../{ef}/")
    print(f"\n👉 Next: python training_pipeline/build_training_data.py\n")

if __name__ == '__main__':
    main()
