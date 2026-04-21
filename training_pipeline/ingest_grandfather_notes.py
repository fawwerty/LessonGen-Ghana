import os
import json
import fitz # PyMuPDF
from docx import Document
from pathlib import Path

ROOT = Path(__file__).parent.parent
NOTES_DIR = ROOT / 'grandfather_notes'
OUTPUT_FILE = ROOT / 'training_pipeline' / 'grandfather_dataset.jsonl'

def extract_text(file_path):
    suffix = file_path.suffix.lower()
    text = ""
    if suffix == '.pdf':
        doc = fitz.open(str(file_path))
        for page in doc:
            text += page.get_text()
        doc.close()
    elif suffix == '.docx':
        doc = Document(str(file_path))
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    return text

def structure_as_example(text, filename):
    # This is a heuristic to guess subject/class from filename
    # We will refine this if possible
    return {
        "messages": [
            {"role": "user", "content": f"Reference example from {filename}"},
            {"role": "assistant", "content": text[:5000]} # Limit size for context
        ]
    }

def main():
    if not NOTES_DIR.exists():
        print("grandfather_notes directory not found")
        return

    examples = []
    for f in NOTES_DIR.iterdir():
        if f.suffix.lower() in ['.pdf', '.docx']:
            print(f"Processing {f.name}...")
            try:
                text = extract_text(f)
                if text.strip():
                    examples.append(structure_as_example(text, f.name))
            except Exception as e:
                print(f"Error processing {f.name}: {e}")

    with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
        for ex in examples:
            f.write(json.dumps(ex, ensure_ascii=False) + '\n')
    
    print(f"Successfully ingested {len(examples)} files into {OUTPUT_FILE}")

if __name__ == "__main__":
    main()
